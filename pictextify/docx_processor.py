#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word文档处理模块 - 负责从Word文档中提取文本和图片
"""

import os
import logging
import docx
from docx.document import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# 获取日志记录器
logger = logging.getLogger("pictextify.docx_processor")


class DocxProcessor:
    """Word文档处理类，负责从Word文档中提取文本和图片"""

    def __init__(self, output_dir=None):
        """
        初始化Word文档处理器

        Args:
            output_dir (str, optional): 图片输出目录
        """
        self.output_dir = output_dir
        logger.debug(f"初始化Word文档处理器: output_dir={output_dir}")

    def process(self, docx_path, output_dir=None):
        """
        处理Word文档，提取文本和图片，保持原始顺序

        Args:
            docx_path (str): Word文档路径
            output_dir (str, optional): 图片输出目录，如果为None则使用初始化时的目录

        Returns:
            tuple: (内容列表, 图片信息列表)
                内容列表: 每个元素为(顺序索引, 内容类型, 内容)
                    内容类型可以是 'text' 或 'image'
                    如果是 'text'，内容为文本字符串
                    如果是 'image'，内容为图片路径
                图片信息列表: 每个元素为(顺序索引, 图片路径)
        """
        # 使用传入的output_dir或初始化时的output_dir
        output_dir = output_dir or self.output_dir
        if not output_dir:
            raise ValueError("必须提供图片输出目录")

        logger.info(f"开始处理Word文档: {docx_path}")
        logger.debug(f"图片输出目录: {output_dir}")

        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)

        try:
            # 打开Word文档
            doc = docx.Document(docx_path)
            logger.info(
                f"Word文档已打开，包含 {len(doc.paragraphs)} 个段落, {len(doc.tables)} 个表格"
            )

            # 提取有序内容
            ordered_items = []
            images_info = []
            item_index = 0

            # 保存所有图片并创建映射
            image_map = {}
            logger.debug("开始提取文档中的图片关系")
            for rel_id, rel in doc.part.rels.items():
                if rel.reltype == RT.IMAGE:
                    # 获取图片数据
                    image_data = rel.target_part.blob

                    # 生成图片文件名
                    image_filename = f"image_{rel_id}.png"
                    image_path = os.path.join(output_dir, image_filename)

                    # 保存图片
                    with open(image_path, "wb") as f:
                        f.write(image_data)

                    # 添加到映射
                    image_map[rel_id] = image_path
                    logger.debug(f"提取图片: rel_id={rel_id}, 保存路径={image_path}")

            logger.info(f"从文档中提取了 {len(image_map)} 张图片")

            # 按顺序处理文档内容
            logger.debug("开始按顺序处理文档内容")

            # 1. 首先处理所有段落，检查每个段落是否包含图片
            for para_idx, paragraph in enumerate(doc.paragraphs):
                # 检查段落是否包含图片
                has_image = False
                img_rel_ids = []

                logger.debug(f"处理段落 {para_idx+1}: 文本长度={len(paragraph.text)}")

                # 检查段落中的所有run
                for run_idx, run in enumerate(paragraph.runs):
                    logger.debug(f"  检查段落 {para_idx+1} 中的run {run_idx+1}")

                    # 使用更简单的方法检查图片
                    if hasattr(run._element, "xpath"):
                        try:
                            # 尝试使用不带namespaces参数的xpath
                            logger.debug(f"  尝试使用xpath查找图片")
                            drawings = run._element.xpath(".//a:blip")

                            # 处理找到的drawings
                            for drawing_idx, drawing in enumerate(drawings):
                                # 获取图片的rel_id
                                rel_id = drawing.get(
                                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                )
                                logger.debug(f"  找到图片引用: rel_id={rel_id}")

                                if rel_id and rel_id in image_map:
                                    has_image = True
                                    img_rel_ids.append(rel_id)
                                    logger.debug(
                                        f"  图片rel_id={rel_id}在映射中找到: {image_map[rel_id]}"
                                    )
                                else:
                                    logger.warning(
                                        f"  图片rel_id={rel_id}在映射中未找到"
                                    )

                        except TypeError:
                            # 如果上面的方法失败，尝试使用字符串搜索方法
                            logger.debug(f"  xpath方法失败，尝试使用XML字符串搜索")
                            try:
                                xml_str = run._element.xml
                                logger.debug(f"  获取到XML字符串，长度={len(xml_str)}")

                                if b"a:blip" in xml_str and b"r:embed" in xml_str:
                                    logger.debug(f"  在XML中找到图片引用")
                                    # 简单解析XML字符串查找r:embed属性值
                                    start_idx = xml_str.find(b'r:embed="') + 9
                                    if start_idx > 8:  # 确保找到了r:embed
                                        end_idx = xml_str.find(b'"', start_idx)
                                        if end_idx > start_idx:
                                            rel_id = xml_str[start_idx:end_idx].decode(
                                                "utf-8"
                                            )
                                            logger.debug(f"  解析出rel_id={rel_id}")

                                            if rel_id in image_map:
                                                has_image = True
                                                img_rel_ids.append(rel_id)
                                                logger.debug(
                                                    f"  图片rel_id={rel_id}在映射中找到: {image_map[rel_id]}"
                                                )
                                            else:
                                                logger.warning(
                                                    f"  图片rel_id={rel_id}在映射中未找到"
                                                )
                                else:
                                    logger.debug(f"  XML中未找到图片引用")
                            except Exception as e:
                                logger.error(f"  解析XML时出错: {str(e)}")
                            continue

                # 添加段落文本
                if paragraph.text.strip():
                    ordered_items.append((item_index, "text", paragraph.text))
                    logger.debug(
                        f"添加文本: 索引={item_index}, 内容长度={len(paragraph.text)}"
                    )
                    item_index += 1
                else:
                    logger.debug(f"段落 {para_idx+1} 文本为空，跳过")

                # 添加段落中的图片
                for rel_id_idx, rel_id in enumerate(img_rel_ids):
                    ordered_items.append((item_index, "image", image_map[rel_id]))
                    images_info.append((item_index, image_map[rel_id]))
                    logger.debug(
                        f"添加图片: 索引={item_index}, 路径={image_map[rel_id]}"
                    )
                    item_index += 1

            # 2. 处理表格
            logger.debug(f"开始处理 {len(doc.tables)} 个表格")
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                logger.debug(
                    f"处理表格 {table_idx+1}: {len(table.rows)}行 x {len(table.rows[0].cells) if table.rows else 0}列"
                )

                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            if para.text.strip():
                                table_text.append(para.text)
                                logger.debug(
                                    f"  表格单元格文本: 行={row_idx+1}, 列={cell_idx+1}, 段落={para_idx+1}, 长度={len(para.text)}"
                                )

                if table_text:
                    combined_text = "\n".join(table_text)
                    ordered_items.append((item_index, "text", combined_text))
                    logger.debug(
                        f"添加表格文本: 索引={item_index}, 内容长度={len(combined_text)}"
                    )
                    item_index += 1
                else:
                    logger.debug(f"表格 {table_idx+1} 文本为空，跳过")

            # 检查是否有未处理的图片（可能是浮动图片或其他特殊情况）
            processed_images = set()
            for _, content_type, content in ordered_items:
                if content_type == "image":
                    processed_images.add(content)

            # 添加未处理的图片
            unprocessed_count = 0
            for rel_id, image_path in image_map.items():
                if image_path not in processed_images:
                    ordered_items.append((item_index, "image", image_path))
                    images_info.append((item_index, image_path))
                    logger.debug(
                        f"添加未处理图片: 索引={item_index}, 路径={image_path}"
                    )
                    item_index += 1
                    unprocessed_count += 1

            if unprocessed_count > 0:
                logger.info(f"添加了 {unprocessed_count} 张未在文本中引用的图片")

            # 按索引排序
            ordered_items.sort(key=lambda x: x[0])
            images_info.sort(key=lambda x: x[0])
            logger.debug("内容已按索引排序")

            logger.info(
                f"Word处理完成: 提取了 {len(ordered_items)} 个内容项, {len(images_info)} 张图片"
            )

            # 打印内容顺序信息，帮助调试
            logger.debug("内容顺序信息:")
            for idx, content_type, content in ordered_items:
                if content_type == "text":
                    logger.debug(
                        f"  索引 {idx}: 文本, 长度={len(content)}, 前20字符: {content[:20]}..."
                    )
                else:
                    logger.debug(f"  索引 {idx}: 图片, 路径={content}")

            return ordered_items, images_info

        except Exception as e:
            logger.exception(f"处理Word文档时出错: {str(e)}")
            raise
